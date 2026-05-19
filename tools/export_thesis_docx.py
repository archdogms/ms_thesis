#!/usr/bin/env python
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def strip_inline(text: str) -> str:
    text = re.sub(r"<sup>\[?([0-9,\-]+)\]?</sup>", r"[\1]", text)
    text = re.sub(r"</?[^>]+>", "", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    return text.strip()


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_indent: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_indent:
        fmt.first_line_indent = Pt(24)


def add_text_paragraph(doc: Document, text: str, *, style: str | None = None, bold: bool = False, center: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_format(paragraph, first_indent=(style is None and not center))
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(strip_inline(text))
    set_run_font(run, 12, bold)


def table_widths_for(title: str | None, max_cols: int) -> list[float] | None:
    if not title:
        return None
    if "13,512 条 POI" in title and max_cols == 13:
        return [0.70] + [1.10] * 12
    if max_cols != 7:
        return None
    if "权重敏感性检验" in title:
        return [1.45, 5.05, 1.35, 1.35, 1.35, 1.35, 1.75]
    if "评论文化识别度" in title:
        return [2.25, 1.15, 1.35, 0.75, 1.25, 1.15, 7.30]
    return None


def apply_table_widths(table, widths_cm: list[float] | None) -> None:
    if not widths_cm:
        table.autofit = True
        return
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_twips = sum(int(Cm(width).twips) for width in widths_cm)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width).twips)))
        grid.append(col)

    for col_idx, width in enumerate(widths_cm):
        table.columns[col_idx].width = Cm(width)
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.first_child_found_in("w:noWrap") is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def add_heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(18 if level == 1 else 12)
    fmt.space_after = Pt(8 if level == 1 else 4)
    run = paragraph.add_run(strip_inline(text))
    set_run_font(run, {1: 16, 2: 14, 3: 13}.get(level, 12), True)


def is_landscape_table(title: str | None) -> bool:
    return False


def set_section_geometry(section, *, landscape: bool) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape and section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    if not landscape and section.page_width > section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)


def add_markdown_table(doc: Document, lines: list[str], title: str | None = None) -> None:
    landscape = is_landscape_table(title)
    if landscape:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(section, landscape=True)
        add_text_paragraph(doc, title or "", bold=True, center=True)

    rows = []
    for i, line in enumerate(lines):
        cells = [strip_inline(cell) for cell in line.strip().strip("|").split("|")]
        if i == 1 and all(re.fullmatch(r":?-+:?", c.strip()) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    widths_cm = table_widths_for(title, max_cols)
    apply_table_widths(table, widths_cm)
    font_size = (
        7.0 if title and "13,512 条 POI" in title
        else 7.0 if title and "评论文化识别度" in title
        else 7.5 if title and "权重敏感性检验" in title
        else 10.5
    )
    nowrap = False
    for r, row in enumerate(rows):
        for c in range(max_cols):
            cell_text = row[c] if c < len(row) else ""
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if nowrap:
                set_cell_no_wrap(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title and c in (1, 6) and "13,512 条 POI" not in title:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(cell_text)
            set_run_font(run, font_size, bold=(r == 0))
    doc.add_paragraph()
    if landscape:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(section, landscape=False)


def resolve_image(md_path: Path, image_ref: str) -> Path:
    ref = image_ref.strip()
    if ref.startswith("<") and ref.endswith(">"):
        ref = ref[1:-1]
    path = Path(ref)
    if path.is_absolute():
        return path
    return (md_path.parent / path).resolve()


def add_image(doc: Document, md_path: Path, image_ref: str) -> None:
    image_path = resolve_image(md_path, image_ref)
    if not image_path.exists():
        add_text_paragraph(doc, f"【图片缺失：{image_ref}】", bold=True, center=True)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width = Inches(4.2) if image_path.name in {"fig_4_2a_kg_global.png", "fig_4_2b_person_network.png"} else Inches(6.2)
    run.add_picture(str(image_path), width=width)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_kg_figure_grid(doc: Document, md_path: Path) -> None:
    top_images = [
        ("fig_4_2a_kg_global.png", "图 4.2　知识图谱全局网络"),
        ("fig_4_2b_person_network.png", "图 4.3　人物关联总图"),
    ]
    figure_dir = md_path.parent / f"{md_path.stem}_media" / "media"
    table = doc.add_table(rows=2, cols=2)
    remove_table_borders(table)
    apply_table_widths(table, [7.5, 7.5])
    for col_idx, (name, caption) in enumerate(top_images):
        path = figure_dir / name
        image_para = table.cell(0, col_idx).paragraphs[0]
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            image_para.add_run().add_picture(str(path), width=Inches(2.9))
        else:
            image_para.add_run(f"【图片缺失：{name}】")
        caption_para = table.cell(1, col_idx).paragraphs[0]
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        set_run_font(run, 10.5)

    huang = figure_dir / "fig_4_3b_huang_feihong.png"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if huang.exists():
        paragraph.add_run().add_picture(str(huang), width=Inches(6.0))
    else:
        paragraph.add_run("【图片缺失：fig_4_3b_huang_feihong.png】")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("图 4.4　黄飞鸿周边人物子图")
    set_run_font(run, 10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    set_section_geometry(section, landscape=False)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for level, size in [(1, 16), (2, 14), (3, 13), (4, 12)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True


def export(md_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_comment = False
    in_code = False
    in_equation = False
    table_buf: list[str] = []
    pending_table_title: str | None = None
    skip_image_names: set[str] = set()
    skip_caption_prefixes: set[str] = set()

    def flush_table() -> None:
        nonlocal table_buf, pending_table_title
        if table_buf:
            add_markdown_table(doc, table_buf, pending_table_title)
            table_buf = []
            pending_table_title = None

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("<!--"):
            in_comment = True
        if in_comment:
            if stripped.endswith("-->"):
                in_comment = False
            continue

        if stripped.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            add_text_paragraph(doc, stripped)
            continue

        if stripped == "$$":
            flush_table()
            in_equation = not in_equation
            continue
        if in_equation:
            add_text_paragraph(doc, stripped, center=True)
            continue

        if stripped.startswith("|"):
            table_buf.append(stripped)
            continue
        flush_table()

        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("综合论文训练记录表"):
            break
        if stripped == "---":
            doc.add_page_break()
            continue

        image = re.match(r"!\[[^\]]*\]\((.*?)\)", stripped)
        if image:
            image_name = resolve_image(md_path, image.group(1)).name
            if image_name in skip_image_names:
                continue
            if image_name == "fig_4_2a_kg_global.png":
                add_kg_figure_grid(doc, md_path)
                skip_image_names.update({"fig_4_2b_person_network.png", "fig_4_3b_huang_feihong.png"})
                skip_caption_prefixes.update({"图 4.2", "图 4.3", "图 4.4"})
                continue
            add_image(doc, md_path, image.group(1))
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            add_heading(doc, len(heading.group(1)), heading.group(2))
            continue

        if re.match(r"^表\s+\d+(?:\.\d+)?", stripped):
            pending_table_title = stripped
            if not is_landscape_table(stripped):
                add_text_paragraph(doc, stripped, bold=True, center=True)
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            set_paragraph_format(paragraph)
            run = paragraph.add_run(strip_inline(stripped[2:]))
            set_run_font(run, 12)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph_format(paragraph)
            run = paragraph.add_run(strip_inline(re.sub(r"^\d+\.\s+", "", stripped)))
            set_run_font(run, 12)
            continue

        only_bold = re.fullmatch(r"\*\*(.*?)\*\*", stripped)
        if only_bold:
            caption_text = strip_inline(only_bold.group(1))
            if any(caption_text.startswith(prefix) for prefix in skip_caption_prefixes):
                continue
            add_text_paragraph(doc, only_bold.group(1), bold=True, center=True)
            continue

        add_text_paragraph(doc, stripped)

    flush_table()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Export the thesis Markdown draft to a Word document.")
    parser.add_argument("--input", required=True, help="Input Markdown path")
    parser.add_argument("--output", required=True, help="Output docx path")
    args = parser.parse_args()

    md_path = (ROOT / args.input).resolve() if not Path(args.input).is_absolute() else Path(args.input)
    output_path = (ROOT / args.output).resolve() if not Path(args.output).is_absolute() else Path(args.output)
    export(md_path, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
